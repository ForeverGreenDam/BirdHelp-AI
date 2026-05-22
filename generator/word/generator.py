"""Word 文件生成器 — 基于 python-docx 将结构化内容构建为 .docx 文件。"""

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from loguru import logger

from generator.base import BaseGenerator

STYLE_THEMES = {
    "academic": {
        "title_color": RGBColor(0x1A, 0x3C, 0x6E),
        "body_color": RGBColor(0x2D, 0x2D, 0x2D),
        "accent_color": RGBColor(0x1A, 0x3C, 0x6E),
        "title_font": "WenQuanYi Micro Hei",
        "body_font": "AR PL UMing CN",
        "heading_font": "WenQuanYi Micro Hei",
    },
    "business": {
        "title_color": RGBColor(0x1B, 0x3A, 0x5C),
        "body_color": RGBColor(0x33, 0x33, 0x33),
        "accent_color": RGBColor(0x00, 0x6E, 0xB6),
        "title_font": "WenQuanYi Micro Hei",
        "body_font": "WenQuanYi Micro Hei",
        "heading_font": "WenQuanYi Micro Hei",
    },
    "creative": {
        "title_color": RGBColor(0xE0, 0x4A, 0x36),
        "body_color": RGBColor(0x3C, 0x3C, 0x3C),
        "accent_color": RGBColor(0xE0, 0x4A, 0x36),
        "title_font": "WenQuanYi Micro Hei",
        "body_font": "WenQuanYi Micro Hei",
        "heading_font": "WenQuanYi Micro Hei",
    },
}


class WordGenerator(BaseGenerator):
    """Word 生成器，将 LLM 输出的结构化 JSON 内容渲染为 .docx 文件。"""

    output_extension = ".docx"

    def generate(self, content: dict[str, Any], output_path: Path) -> Path:
        """根据结构化内容生成 Word 文件，返回输出路径。"""
        parsed = self._parse_content(content)
        title = parsed.get("title", "文档")
        subtitle = parsed.get("subtitle", "")
        abstract = parsed.get("abstract", "")
        sections = parsed.get("sections", [])
        references = parsed.get("references", [])
        style_name = parsed.get("style", "academic")
        theme = STYLE_THEMES.get(style_name, STYLE_THEMES["academic"])

        doc = Document()

        # 页面设置：A4
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # 标题页
        self._add_title_page(doc, title, subtitle, theme)

        # 摘要
        if abstract:
            self._add_abstract(doc, abstract, theme)

        # 章节
        for section_data in sections:
            heading = section_data.get("heading", "")
            content_items = section_data.get("content", [])
            if heading:
                self._add_heading(doc, heading, 1, theme)
            for para_text in content_items:
                self._add_paragraph(doc, para_text, theme, "body")

        # 参考文献
        if references:
            self._add_references(doc, references, theme)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(f"Word generated: {output_path}, {len(sections)} sections, style={style_name}")
        return output_path

    def _add_title_page(self, doc: Document, title: str, subtitle: str, theme: dict) -> None:
        """添加标题页。"""
        # 空行留白
        for _ in range(6):
            doc.add_paragraph("")

        # 主标题
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = theme["title_color"]
        run.font.name = theme["title_font"]
        run.element.rPr.rFonts.set(qn("w:eastAsia"), theme["title_font"])

        if subtitle:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(subtitle)
            run.font.size = Pt(16)
            run.font.color.rgb = theme["accent_color"]
            run.font.name = theme["body_font"]
            run.element.rPr.rFonts.set(qn("w:eastAsia"), theme["body_font"])

        # 日期
        from datetime import date
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(str(date.today()))
        run.font.size = Pt(12)
        run.font.color.rgb = theme["body_color"]
        run.font.name = theme["body_font"]
        run.element.rPr.rFonts.set(qn("w:eastAsia"), theme["body_font"])

        # 分页
        doc.add_page_break()

    def _add_abstract(self, doc: Document, abstract: str, theme: dict) -> None:
        """添加摘要部分。"""
        self._add_heading(doc, "摘要", 1, theme)
        self._add_paragraph(doc, abstract, theme, "body")
        doc.add_page_break()

    def _add_references(self, doc: Document, references: list[str], theme: dict) -> None:
        """添加参考文献部分。"""
        self._add_heading(doc, "参考文献", 1, theme)
        for i, ref in enumerate(references, 1):
            para = doc.add_paragraph()
            run = para.add_run(f"[{i}] {ref}")
            run.font.size = Pt(10.5)
            run.font.color.rgb = theme["body_color"]
            run.font.name = theme["body_font"]
            run.element.rPr.rFonts.set(qn("w:eastAsia"), theme["body_font"])
            para.paragraph_format.space_after = Pt(4)

    def _add_heading(self, doc: Document, text: str, level: int, theme: dict) -> None:
        """添加格式化标题。"""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = theme["title_color"]
            run.font.name = theme["heading_font"]
            run.element.rPr.rFonts.set(qn("w:eastAsia"), theme["heading_font"])

    def _add_paragraph(self, doc: Document, text: str, theme: dict, role: str = "body") -> None:
        """添加格式化段落。"""
        para = doc.add_paragraph()
        self._style_paragraph(para, theme, role)
        run = para.add_run(text)
        if role == "body":
            run.font.size = Pt(12)
            run.font.color.rgb = theme["body_color"]
        run.font.name = theme["body_font"]
        run.element.rPr.rFonts.set(qn("w:eastAsia"), theme["body_font"])
        para.paragraph_format.first_line_indent = Inches(0.28)

    @staticmethod
    def _style_paragraph(para, theme: dict, role: str) -> None:
        """根据角色设置段落样式。"""
        if role == "body":
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
