"""python-docx 文档构建器 — Word 和 PDF 共用的增强型文档渲染逻辑。

支持封面、标题层级、段落、图表（PNG嵌入）、图片、表格、分栏、页眉页脚。
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from generator._design import ColorPalette
from generator._chart_engine import render_chart
from utils.file import ensure_temp_dir


def hex_to_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_shading(cell, color_hex: str) -> None:
    """设置表格单元格背景色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex.lstrip("#"))
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def _set_run_font(run, font_name: str, size_pt: int, color: RGBColor,
                  bold: bool = False) -> None:
    """统一样式设置 helper — 同时设置拉丁和东亚字体。"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _add_styled_paragraph(doc: Document, text: str, palette: ColorPalette,
                          font_size: int = 12, bold: bool = False,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT,
                          first_line_indent: float | None = 0.28,
                          space_after: int = 6, line_spacing: float = 1.5) -> None:
    """添加段落并应用样式。"""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = line_spacing
    if first_line_indent is not None:
        para.paragraph_format.first_line_indent = Inches(first_line_indent)
    run = para.add_run(text)
    _set_run_font(run, palette.body_font, font_size, hex_to_rgb(palette.text_body), bold)


class DocxBuilder:
    """Word/PDF 文档构建器 — 将 LLM 输出的结构化 JSON 渲染为 python-docx Document。"""

    def __init__(self, palette: ColorPalette, enable_images: bool = True,
                 enable_charts: bool = True):
        self.palette = palette
        self.enable_images = enable_images
        self.enable_charts = enable_charts
        self._chart_dir: Path | None = None
        self._image_dir: Path | None = None

    @property
    def chart_dir(self) -> Path:
        if self._chart_dir is None:
            self._chart_dir = ensure_temp_dir() / "charts"
            self._chart_dir.mkdir(parents=True, exist_ok=True)
        return self._chart_dir

    @property
    def image_dir(self) -> Path:
        if self._image_dir is None:
            self._image_dir = ensure_temp_dir() / "doc_images"
            self._image_dir.mkdir(parents=True, exist_ok=True)
        return self._image_dir

    # ── 顶层入口 ──

    def build_document(self, parsed: dict) -> Document:
        """将解析后的 JSON 构建为完整的 python-docx Document。"""
        doc = Document()
        self._setup_page(doc)

        title = parsed.get("title", "文档")
        subtitle = parsed.get("subtitle", "")
        abstract = parsed.get("abstract", "")
        sections = parsed.get("sections", [])
        tables = parsed.get("tables", [])
        references = parsed.get("references", [])
        author = parsed.get("author", "")
        doc_date = parsed.get("date", "")

        # 封面
        self.add_cover(doc, title, subtitle, author, doc_date)

        # 摘要
        if abstract:
            self.add_abstract(doc, abstract)

        # 章节（含内嵌 charts/images）
        for sec in sections:
            self.add_section(doc, sec)

        # 文档级表格
        for tbl in tables:
            self.add_table(doc, tbl)

        # 参考文献
        if references:
            self.add_references(doc, references)

        return doc

    # ── 页面设置 ──

    def _setup_page(self, doc: Document) -> None:
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── 封面 ──

    def add_cover(self, doc: Document, title: str, subtitle: str,
                  author: str = "", doc_date: str = "") -> None:
        p = self.palette
        # 顶部色条
        top_para = doc.add_paragraph()
        top_para.paragraph_format.space_after = Pt(0)
        # 留白
        for _ in range(5):
            doc.add_paragraph("")

        # 标题
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        _set_run_font(run, p.title_font, 28, hex_to_rgb(p.primary), bold=True)

        if subtitle:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(subtitle)
            _set_run_font(run, p.body_font, 16, hex_to_rgb(p.accent))

        # 元信息
        meta_parts = [p for p in [author, doc_date or str(date.today())] if p]
        if meta_parts:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(" | ".join(meta_parts))
            _set_run_font(run, p.body_font, 12, hex_to_rgb(p.text_secondary))

        doc.add_page_break()

    # ── 摘要 ──

    def add_abstract(self, doc: Document, text: str) -> None:
        self._add_heading(doc, "摘要", 1)
        _add_styled_paragraph(doc, text, self.palette, font_size=12,
                              first_line_indent=0.28)
        doc.add_page_break()

    # ── 章节（含 chart / image 内嵌） ──

    def add_section(self, doc: Document, section_data: dict) -> None:
        heading = section_data.get("heading", "")
        content_items = section_data.get("content", [])
        charts = section_data.get("charts", [])
        images = section_data.get("images", [])

        if heading:
            self._add_heading(doc, heading, 1)

        for para_text in content_items:
            _add_styled_paragraph(doc, para_text, self.palette, font_size=12,
                                  first_line_indent=0.28)

        # 图表
        if self.enable_charts:
            for chart_spec in charts:
                chart_path = self.chart_dir / f"chart_{hash(str(chart_spec)) & 0x7FFFFFFF:x}.png"
                result = render_chart(chart_spec, chart_path, self.palette)
                if result and result.exists():
                    self._insert_image(doc, str(result), chart_spec.get("width", "full"),
                                       WD_ALIGN_PARAGRAPH.CENTER)
                    caption = chart_spec.get("caption", "")
                    if caption:
                        _add_styled_paragraph(doc, caption, self.palette, font_size=10,
                                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                              first_line_indent=None, space_after=12)

        # 图片
        if self.enable_images:
            for img_spec in images:
                img_path = img_spec.get("_local_path", "")
                if img_path and Path(img_path).exists():
                    self._insert_image(doc, img_path, img_spec.get("width", "half"),
                                       WD_ALIGN_PARAGRAPH.CENTER)
                    caption = img_spec.get("caption", "")
                    if caption:
                        _add_styled_paragraph(doc, caption, self.palette, font_size=10,
                                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                              first_line_indent=None, space_after=12)

    # ── 表格 ──

    def add_table(self, doc: Document, table_data: dict) -> None:
        p = self.palette
        caption = table_data.get("caption", "")
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        if not headers:
            return

        if caption:
            _add_styled_paragraph(doc, caption, p, font_size=12, bold=True,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  first_line_indent=None, space_after=8)

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"

        # 表头
        header_color = hex_to_rgb(p.primary)
        header_bg = p.primary
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            _set_cell_shading(cell, header_bg)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    _set_run_font(run, p.body_font, 11, RGBColor(0xFF, 0xFF, 0xFF), bold=True)

        # 数据行（斑马条纹）
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = str(cell_text)
                if row_idx % 2 == 1:
                    _set_cell_shading(cell, p.light)
                for para in cell.paragraphs:
                    for run in para.runs:
                        _set_run_font(run, p.body_font, 11, hex_to_rgb(p.text_body))

        doc.add_paragraph("")

    # ── 参考文献 ──

    def add_references(self, doc: Document, references: list[str]) -> None:
        self._add_heading(doc, "参考文献", 1)
        for i, ref in enumerate(references, 1):
            _add_styled_paragraph(doc, f"[{i}] {ref}", self.palette, font_size=10.5,
                                  first_line_indent=None, space_after=4, line_spacing=1.2)

    # ── 内部工具 ──

    def _add_heading(self, doc: Document, text: str, level: int) -> None:
        p = self.palette
        heading = doc.add_heading(text, level=level)
        sizes = {1: 20, 2: 16, 3: 14}
        for run in heading.runs:
            _set_run_font(run, p.heading_font, sizes.get(level, 14),
                          hex_to_rgb(p.primary), bold=True)

    def _insert_image(self, doc: Document, path: str, width: str,
                      alignment) -> None:
        """插入图片，最大宽度限制避免溢出。"""
        try:
            para = doc.add_paragraph()
            para.alignment = alignment
            run = para.add_run()
            width_inches = {"full": 5.5, "half": 3.0, "quarter": 1.5}.get(width, 3.0)
            run.add_picture(path, width=Inches(min(width_inches, 5.5)))
            doc.add_paragraph("")
        except Exception:
            pass
