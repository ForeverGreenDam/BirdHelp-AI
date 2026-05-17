"""PDF 文件生成器 — 基于 python-docx + LibreOffice 将结构化内容构建为 .pdf 文件。"""

import shutil
import subprocess
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger

from core.exceptions import FileGenerationError
from generator.base import BaseGenerator
from utils.file import temp_file_path, ensure_temp_dir

STYLE_THEMES = {
    "academic": {
        "title_color": RGBColor(0x1A, 0x3C, 0x6E),
        "body_color": RGBColor(0x2D, 0x2D, 0x2D),
        "accent_color": RGBColor(0x1A, 0x3C, 0x6E),
        "title_font": "WenQuanYi Micro Hei",
        "body_font": "AR PL UMing CN",
        "heading_font": "WenQuanYi Micro Hei",
    },
    "business": {
        "title_color": RGBColor(0x1B, 0x3A, 0x5C),
        "body_color": RGBColor(0x33, 0x33, 0x33),
        "accent_color": RGBColor(0x00, 0x6E, 0xB6),
        "title_font": "WenQuanYi Micro Hei",
        "body_font": "WenQuanYi Micro Hei",
        "heading_font": "WenQuanYi Micro Hei",
    },
    "creative": {
        "title_color": RGBColor(0xE0, 0x4A, 0x36),
        "body_color": RGBColor(0x3C, 0x3C, 0x3C),
        "accent_color": RGBColor(0xE0, 0x4A, 0x36),
        "title_font": "WenQuanYi Micro Hei",
        "body_font": "WenQuanYi Micro Hei",
        "heading_font": "WenQuanYi Micro Hei",
    },
}


class PdfGenerator(BaseGenerator):
    """PDF 生成器，将 LLM 输出的结构化 JSON 文档内容先构建为 .docx，再转换为 .pdf 文件。"""

    output_extension = ".pdf"

    def generate(self, content: dict[str, Any], output_path: Path) -> Path:
        """根据结构化内容生成 PDF 文件，返回输出路径。"""
        parsed = self._parse_content(content)
        title = parsed.get("title", "文档")
        style_name = parsed.get("style", "academic")
        theme = STYLE_THEMES.get(style_name, STYLE_THEMES["academic"])

        # 1. 创建临时 .docx
        docx_path = temp_file_path(".docx")
        doc = Document()

        # 页面设置: A4, 1 英寸边距
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # 2. 构建文档内容
        subtitle = parsed.get("subtitle", "")
        author = parsed.get("author", "")
        date = parsed.get("date", "")

        # 标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = theme["title_color"]
        run.font.name = theme["title_font"]

        # 副标题
        if subtitle:
            sub_para = doc.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sub_para.add_run(subtitle)
            run.font.size = Pt(16)
            run.font.color.rgb = theme["body_color"]
            run.font.name = theme["body_font"]

        # 作者和日期
        info_parts = [p for p in [author, date] if p]
        if info_parts:
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = info_para.add_run(" | ".join(info_parts))
            run.font.size = Pt(12)
            run.font.color.rgb = theme["body_color"]
            run.font.name = theme["body_font"]

        # 分隔线
        doc.add_paragraph()

        # 章节
        sections_data = parsed.get("sections", [])
        for section_data in sections_data:
            self._add_heading(doc, section_data.get("heading", ""), 1, theme)
            for para_text in section_data.get("content", []):
                self._add_paragraph(doc, para_text, theme, "body")

        # 表格
        tables_data = parsed.get("tables", [])
        for table_data in tables_data:
            self._add_table(
                doc,
                table_data.get("caption", ""),
                table_data.get("headers", []),
                table_data.get("rows", []),
                theme,
            )

        # 3. 保存 .docx
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(docx_path))
        logger.info(f"DOCX generated: {docx_path}")

        # 4. 转换为 PDF
        output_dir = output_path.parent
        output_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = self._convert_to_pdf(docx_path, output_dir)

        if pdf_path:
            if pdf_path != output_path:
                shutil.move(str(pdf_path), str(output_path))
            logger.info(f"PDF generated: {output_path}, style={style_name}")
            if docx_path.exists():
                docx_path.unlink()
            return output_path
        else:
            if docx_path.exists():
                docx_path.unlink()
            raise FileGenerationError(
                "LibreOffice 转换 PDF 失败，请检查服务端 LibreOffice 是否正常安装"
            )

    def _add_heading(self, doc: Document, text: str, level: int, theme: dict) -> None:
        """添加格式化标题。"""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = theme["accent_color"]
            run.font.name = theme["heading_font"]
            run.font.size = Pt(18 if level == 1 else 15)

    def _add_paragraph(self, doc: Document, text: str, theme: dict, role: str = "body") -> None:
        """添加格式化段落。"""
        para = doc.add_paragraph()
        para.add_run(text)
        self._style_paragraph(para, theme, role)

    def _add_table(self, doc: Document, caption: str, headers: list[str], rows: list[list[str]], theme: dict) -> None:
        """添加格式化表格。"""
        if caption:
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_para.add_run(caption)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = theme["accent_color"]
            run.font.name = theme["body_font"]

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"

        # 表头
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = theme["accent_color"]
                    run.font.name = theme["body_font"]

        # 数据行
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.color.rgb = theme["body_color"]
                        run.font.name = theme["body_font"]

        doc.add_paragraph()

    @staticmethod
    def _convert_to_pdf(docx_path: Path, output_dir: Path) -> Optional[Path]:
        """使用 LibreOffice 将 .docx 转换为 .pdf，返回 pdf 路径。"""
        try:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", str(output_dir), str(docx_path)],
                capture_output=True,
                text=True,
                timeout=120,
            )
            if result.returncode != 0:
                logger.error(f"LibreOffice conversion failed: {result.stderr}")
                return None
            logger.info(f"LibreOffice conversion output: {result.stdout}")
            pdf_path = output_dir / docx_path.with_suffix(".pdf").name
            if pdf_path.exists():
                return pdf_path
            return None
        except FileNotFoundError:
            logger.warning("LibreOffice not found on system")
            return None
        except Exception as exc:
            logger.error(f"LibreOffice conversion error: {exc}")
            return None

    @staticmethod
    def _style_paragraph(para, theme: dict, role: str) -> None:
        """根据角色设置段落样式。"""
        if role == "title":
            for run in para.runs:
                run.font.size = Pt(26)
                run.font.bold = True
                run.font.color.rgb = theme["title_color"]
                run.font.name = theme["title_font"]
        elif role == "heading":
            for run in para.runs:
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = theme["accent_color"]
                run.font.name = theme["heading_font"]
        elif role == "body":
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.color.rgb = theme["body_color"]
                run.font.name = theme["body_font"]
